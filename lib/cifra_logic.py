import requests
from bs4 import BeautifulSoup
from fpdf import FPDF
import sys
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class PDF(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 15)
        # Title will be set in the main logic
        
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.cell(0, 10, f'Página {self.page_no()}/{{nb}}', align='C')

NOTES_SHARP = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']
NOTES_FLAT = ['C', 'Db', 'D', 'Eb', 'E', 'F', 'Gb', 'G', 'Ab', 'A', 'Bb', 'B']

def get_note_index(note):
    # Normalize note
    note = note.replace('Cb', 'B').replace('B#', 'C').replace('Fb', 'E').replace('E#', 'F')
    if note in NOTES_SHARP:
        return NOTES_SHARP.index(note)
    if note in NOTES_FLAT:
        return NOTES_FLAT.index(note)
    return -1

def transpose_note(note, semitones, use_flats=False):
    idx = get_note_index(note)
    if idx == -1: return note
    new_idx = (idx + semitones) % 12
    return NOTES_FLAT[new_idx] if use_flats else NOTES_SHARP[new_idx]

def transpose_chord(chord, semitones, use_flats=False):
    # Split root and bass
    parts = chord.split('/')
    root_part = parts[0]
    bass_part = parts[1] if len(parts) > 1 else None
    
    # Extract root note
    root_len = 0
    if len(root_part) > 1 and root_part[1] in ['#', 'b']:
        root_len = 2
    else:
        root_len = 1
    
    root_note = root_part[:root_len]
    suffix = root_part[root_len:]
    
    new_root = transpose_note(root_note, semitones, use_flats)
    new_chord = new_root + suffix
    
    if bass_part:
        bass_len = 0
        if len(bass_part) > 1 and bass_part[1] in ['#', 'b']:
            bass_len = 2
        else:
            bass_len = 1
        bass_note = bass_part[:bass_len]
        bass_suffix = bass_part[bass_len:]
        
        new_bass = transpose_note(bass_note, semitones, use_flats)
        new_chord += '/' + new_bass + bass_suffix
        
    return new_chord

def transpose_lines(lines, semitones, use_flats=False):
    new_lines = []
    for line in lines:
        if is_chord_line(line):
            new_line = []
            for segment in line:
                if segment.get('bold'): # It's a chord
                    text = segment['text']
                    parts = text.split(' ')
                    new_parts = []
                    for part in parts:
                        if not part:
                            new_parts.append('')
                            continue
                        # Simple check if part starts with note
                        if part[0] in 'ABCDEFG':
                            new_parts.append(transpose_chord(part, semitones, use_flats))
                        else:
                            new_parts.append(part)
                    new_text = ' '.join(new_parts)
                    
                    new_segment = segment.copy()
                    new_segment['text'] = new_text
                    new_line.append(new_segment)
                else:
                    new_line.append(segment)
            new_lines.append(new_line)
        else:
            new_lines.append(line)
    return new_lines

def get_cifra_content(url, target_key_index=None):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise Exception(f"Erro ao acessar a URL: {e}")

    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract Title and Artist
    title_tag = soup.find('h1', class_='t1')
    artist_tag = soup.find('h2', class_='t3')
    
    title = title_tag.get_text(strip=True) if title_tag else "Título Desconhecido"
    artist = artist_tag.get_text(strip=True) if artist_tag else "Artista Desconhecido"

    # Extract Key (Tom)
    key_element = soup.find('span', id='cifra_tom')
    key = key_element.get_text(strip=True) if key_element else ""
    if key:
        key = f"Tom: {key}"

    # Extract Cipher Content
    pre_content = soup.find('pre')
    
    if not pre_content:
        raise Exception("Não foi possível encontrar a cifra (tag <pre> não encontrada).")

    # Parse content preserving bold tags
    lines = []
    current_line = []
    
    for element in pre_content.contents:
        if element.name == 'b':
            text = element.get_text()
            current_line.append({'text': text, 'bold': True})
        elif element.name == 'br':
            lines.append(current_line)
            current_line = []
        elif isinstance(element, str) or (element.string and element.name is None):
            text = str(element)
            parts = text.split('\n')
            for i, part in enumerate(parts):
                if part:
                    current_line.append({'text': part, 'bold': False})
                if i < len(parts) - 1:
                    lines.append(current_line)
                    current_line = []
    
    if current_line:
        lines.append(current_line)
    
    # Transposition Logic
    if target_key_index is not None and key:
        try:
            # Clean key text more robustly
            clean_key = key.lower().replace("tom:", "").strip()
            # Restore case for note parsing (first letter upper)
            if clean_key:
                clean_key = clean_key[0].upper() + clean_key[1:]
            
            original_key_note = clean_key
            
            print(f"DEBUG: Scraped key: '{key}', Cleaned: '{original_key_note}'", file=sys.stderr)
            
            # Handle minor keys? "Cm" -> "C"
            is_minor = 'm' in original_key_note
            if is_minor:
                original_key_note = original_key_note.replace('m', '')
            
            original_idx = get_note_index(original_key_note)
            print(f"DEBUG: Original Key Index: {original_idx}, Target (Cifra): {target_key_index}", file=sys.stderr)
            
            if original_idx != -1:
                # Explicit mapping based on Cifra Club values
                # C: key=3, D: key=5, E: key=7, F: key=8, G: key=10, A: key=0, B: key=2
                # We map these to our chromatic index (0=C, 1=C#, etc.)
                CIFRA_CLUB_KEY_MAP = {
                    0: 9,   # A
                    1: 10,  # A# / Bb
                    2: 11,  # B
                    3: 0,   # C
                    4: 1,   # C# / Db
                    5: 2,   # D
                    6: 3,   # D# / Eb
                    7: 4,   # E
                    8: 5,   # F
                    9: 6,   # F# / Gb
                    10: 7,  # G
                    11: 8   # G# / Ab
                }
                
                if target_key_index in CIFRA_CLUB_KEY_MAP:
                    target_chromatic_index = CIFRA_CLUB_KEY_MAP[target_key_index]
                else:
                    # Fallback to formula if key is outside 0-11 range (though unlikely)
                    print(f"WARN: Unknown key index {target_key_index}, using formula.", file=sys.stderr)
                    target_chromatic_index = (target_key_index + 9) % 12
                
                semitones = target_chromatic_index - original_idx
                print(f"DEBUG: Target Chromatic: {target_chromatic_index}, Semitones: {semitones}", file=sys.stderr)
                
                # Determine target key name for display and accidental preference
                # Heuristic for flats: F(5), Bb(10), Eb(3), Ab(8), Db(1)
                use_flats = target_chromatic_index in [1, 3, 5, 8, 10]
                
                lines = transpose_lines(lines, semitones, use_flats)
                
                # Update key text
                new_key_note = NOTES_FLAT[target_chromatic_index] if use_flats else NOTES_SHARP[target_chromatic_index]
                key = f"Tom: {new_key_note}{'m' if is_minor else ''}"
        except Exception as e:
            print(f"Error transposing: {e}", file=sys.stderr)

    return title, artist, key, lines

def is_chord_line(line_segments):
    has_bold = any(s['bold'] for s in line_segments)
    if has_bold:
        return True
        
    text = "".join([s['text'] for s in line_segments]).strip()
    if not text:
        return False
        
    allowed = set("ABCDEFGMmb#0123456789/()+-^°ºdimsusaug ")
    if all(c in allowed for c in text):
        return True
        
    return False

def is_header_line(line_segments):
    text = "".join([s['text'] for s in line_segments]).strip()
    return text.startswith('[') and text.endswith(']')

def deduplicate_sections(lines):
    deduplicated_lines = []
    seen_headers = set()
    skip_mode = False
    
    for line in lines:
        if is_header_line(line):
            text = "".join([s['text'] for s in line]).strip()
            if text in seen_headers:
                skip_mode = True
                new_line = []
                for s in line:
                    new_s = s.copy()
                    new_s['italic'] = True
                    new_line.append(new_s)
                deduplicated_lines.append(new_line)
            else:
                seen_headers.add(text)
                skip_mode = False
                deduplicated_lines.append(line)
        else:
            if not skip_mode:
                deduplicated_lines.append(line)
                
    return deduplicated_lines

def pair_lines(lines):
    units = []
    i = 0
    while i < len(lines):
        line = lines[i]
        text = "".join([s['text'] for s in line]).strip()
        
        if not text:
            i += 1
            continue
            
        if is_header_line(line):
            units.append({'type': 'header', 'chords': [], 'lyrics': line})
            i += 1
            continue
            
        if is_chord_line(line):
            if i + 1 < len(lines):
                next_line = lines[i+1]
                if is_header_line(next_line):
                    units.append({'type': 'pair', 'chords': line, 'lyrics': []})
                    i += 1
                elif not is_chord_line(next_line):
                    units.append({'type': 'pair', 'chords': line, 'lyrics': next_line})
                    i += 2
                else:
                    units.append({'type': 'pair', 'chords': line, 'lyrics': []})
                    i += 1
            else:
                units.append({'type': 'pair', 'chords': line, 'lyrics': []})
                i += 1
        else:
            units.append({'type': 'pair', 'chords': [], 'lyrics': line})
            i += 1
    return units

def get_line_length(line_segments):
    return len("".join([s['text'] for s in line_segments]))

def pad_line(line_segments, target_length):
    current_len = get_line_length(line_segments)
    if current_len < target_length:
        padding = " " * (target_length - current_len)
        if line_segments:
            new_segments = list(line_segments)
            new_segments.append({'text': padding, 'bold': False})
            return new_segments
        else:
            return [{'text': " " * target_length, 'bold': False}]
    return line_segments

def reflow_units(units, max_chars):
    rows = []
    current_row_units = []
    current_row_len = 0
    gap = 3
    
    for i, unit in enumerate(units):
        if unit.get('type') == 'header':
            if current_row_units:
                rows.extend(build_row_lines(current_row_units, gap))
                current_row_units = []
                current_row_len = 0
            
            if rows: 
                rows.append([{'text': ' ', 'bold': False}])

            l_len = get_line_length(unit['lyrics'])
            rows.extend(build_row_lines([(unit, l_len)], 0))
            continue

        c_len = get_line_length(unit['chords'])
        l_len = get_line_length(unit['lyrics'])
        unit_width = max(c_len, l_len)
        
        needed = unit_width
        if current_row_units:
            needed += gap
            
        if current_row_len + needed <= max_chars:
            current_row_units.append((unit, unit_width))
            current_row_len += needed
        else:
            if current_row_units:
                rows.extend(build_row_lines(current_row_units, gap))
            
            current_row_units = [(unit, unit_width)]
            current_row_len = unit_width
            
    if current_row_units:
        rows.extend(build_row_lines(current_row_units, gap))
        
    return rows

def build_row_lines(row_units, gap):
    final_chords = []
    final_lyrics = []
    
    for i, (unit, width) in enumerate(row_units):
        if i > 0:
            gap_seg = [{'text': " " * gap, 'bold': False}]
            final_chords.extend(gap_seg)
            final_lyrics.extend(gap_seg)
            
        padded_chords = pad_line(unit['chords'], width)
        final_chords.extend(padded_chords)
        
        padded_lyrics = pad_line(unit['lyrics'], width)
        final_lyrics.extend(padded_lyrics)
        
    result = []
    if "".join([s['text'] for s in final_chords]).strip():
        result.append(final_chords)
    if "".join([s['text'] for s in final_lyrics]).strip():
        result.append(final_lyrics)
        
    return result

def calculate_layout(lines, available_height, available_width):
    lines = deduplicate_sections(lines)
    options = [12, 11.5, 11, 10.5, 10, 9.5, 9, 8.5, 8]
    units = pair_lines(lines)
    
    for font_size in options:
        line_height = (font_size * 1.2) * 0.3527
        char_width_mm = (font_size * 0.6) * 0.3527
        max_chars = int(available_width / char_width_mm)
        reflowed_lines = reflow_units(units, max_chars)
        total_height_needed = len(reflowed_lines) * line_height
        
        if total_height_needed <= available_height:
            return font_size, reflowed_lines, line_height
            
    fallback_size = 8
    max_chars = int(available_width / ((fallback_size * 0.6) * 0.3527))
    return fallback_size, reflow_units(units, max_chars), (fallback_size * 1.2) * 0.3527

def generate_pdf_bytes(title, artist, key, lines):
    pdf = PDF(orientation='P')
    pdf.set_margins(5, 5, 5)
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Title: Bold + Underline
    pdf.set_font('Helvetica', 'BU', 14)
    # Reduced height from 7 to 5 to bring artist closer
    pdf.cell(0, 5, title, new_x="LMARGIN", new_y="NEXT", align='C')
    
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 7, artist, new_x="LMARGIN", new_y="NEXT", align='C')
    
    # Header height calculation (Key removed)
    header_height = 15 # Adjusted base height
    
    usable_width = 200
    usable_height = 287 - header_height
    
    font_size, reflowed_lines, line_height = calculate_layout(lines, usable_height, usable_width)
    
    pdf.set_font('Courier', '', font_size)
    
    initial_y = pdf.get_y()
    pdf.set_xy(5, initial_y)
    
    for line_segments in reflowed_lines:
        pdf.set_x(5)
        for segment in line_segments:
            text = segment['text']
            is_bold = segment['bold']
            is_italic = segment.get('italic', False)
            
            style = ''
            if is_bold: style += 'B'
            if is_italic: style += 'I'
            
            pdf.set_font('Courier', style, font_size)
            pdf.write(line_height, text)
        pdf.ln(line_height)

    # Return bytes
    return pdf.output(dest='S')

def generate_docx_bytes(title, artist, key, lines):
    usable_width = 200
    usable_height = 270
    font_size, reflowed_lines, line_height = calculate_layout(lines, usable_height, usable_width)
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Remove space after title to bring artist closer
    p_title.paragraph_format.space_after = Pt(0)
    
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.underline = True
    run_title.font.name = 'Helvetica'
    run_title.font.size = Pt(14)

    p_artist = doc.add_paragraph()
    p_artist.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_artist = p_artist.add_run(artist)
    run_artist.font.name = 'Helvetica'
    run_artist.font.size = Pt(10)
    
    # Key (Tom) section removed
    
    font_name = 'Courier New'
    
    font_name = 'Courier New'
    
    for line_segments in reflowed_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        
        for segment in line_segments:
            text = segment['text']
            is_bold = segment['bold']
            is_italic = segment.get('italic', False)
            
            run = p.add_run(text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = is_bold
            run.italic = is_italic

    # Return bytes
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()

def get_content_from_file(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        raise Exception(f"Erro ao ler o arquivo: {e}")

    lines = []
    raw_lines = content.splitlines()
    
    title = "Título Desconhecido"
    artist = "Artista Desconhecido"
    key = ""
    
    start_index = 0
    if len(raw_lines) > 0 and not is_chord_line([{'text': raw_lines[0], 'bold': False}]):
        title = raw_lines[0].strip()
        start_index += 1
    if len(raw_lines) > 1 and not is_chord_line([{'text': raw_lines[1], 'bold': False}]):
        artist = raw_lines[1].strip()
        start_index += 1
        
    for i in range(start_index, min(start_index + 5, len(raw_lines))):
        if "Tom:" in raw_lines[i]:
            key = raw_lines[i].strip()
            break

    for i in range(start_index, len(raw_lines)):
        text = raw_lines[i].rstrip()
        temp_segments = [{'text': text, 'bold': False}]
        
        if is_chord_line(temp_segments):
            lines.append([{'text': text, 'bold': True}])
        else:
            lines.append([{'text': text, 'bold': False}])
            
    return title, artist, key, lines
